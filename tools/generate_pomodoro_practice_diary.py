from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


TEMPLATE_PATH = Path('/Users/kamilhus/Downloads/Дневник студента ФИИТ ПТП 231 Мудрик.docx')
OUTPUT_PATH = Path('/Users/kamilhus/Downloads/Дневник студента ФИИТ ПТП 231 Хусаинов.docx')

STUDENT_FULL_NAME = 'Хусаинов Камиль Фанисович'
GROUP = '09-231'
START_DATE = 'Дата начала практики         «10» февраля  2026 г.'
END_DATE = 'Дата окончания практики   «25» марта 2026 г.'
SUPERVISOR_SHORT = 'Курбангалеев А.А.'
PRACTICE_PLACE = 'КФУ, Институт ВМиИТ, кафедра системного анализа и информационных технологий'

DAILY_TASKS = [
    'Анализ существующих приложений для Pomodoro, трекеров целей и задач',
    'Анализ существующих приложений для Pomodoro, трекеров целей и задач',
    'Формирование пользовательских сценариев и анализ предметной области',
    'Формирование пользовательских сценариев и анализ предметной области',
    'Проектирование общей архитектуры монорепозитория и модулей системы',
    'Проектирование общей архитектуры монорепозитория и модулей системы',
    'Проектирование общей архитектуры монорепозитория и модулей системы',
    'Разработка функциональных требований к приложению Pomodoro Web',
    'Разработка нефункциональных требований и сценариев безопасности',
    'Проектирование пользовательского интерфейса и навигации',
    'Проектирование пользовательского интерфейса и навигации',
    'Проектирование структуры базы данных PostgreSQL',
    'Проектирование структуры базы данных PostgreSQL',
    'Подготовка REST API для целей, задач и фокус-сессий',
    'Реализация доменных сущностей backend и миграций Flyway',
    'Реализация JWT-аутентификации и авторизации пользователей',
    'Реализация API для целей и задач',
    'Реализация API для фокус-сессий и отчетов',
    'Реализация планировщика и расчетов прогресса и streak',
    'Интеграция AI-сервиса для анализа фото-отчетов',
    'Интеграция AI-сервиса для чата-мотиватора',
    'Реализация мотивационной ленты и загрузки изображений',
    'Реализация экранов входа и регистрации на frontend',
    'Реализация вкладки Контроль и управления целями',
    'Реализация вкладки Фокус и таймера Pomodoro',
    'Реализация вкладки Фокус и AI-проверки фото-отчетов',
    'Реализация вкладки Мотивация с лентой изображений и цитат',
    'Реализация вкладки Мотивация с лентой изображений и цитат',
    'Реализация экрана Мотиватор с хранением истории диалога',
    'Реализация страницы статистики и визуализации прогресса',
    'Интеграция frontend и backend, настройка Docker Compose',
    'Написание и запуск backend- и frontend-тестов',
    'Подготовка отчета и дневника по практике',
    'Подготовка отчета и дневника по практике',
    'Финальная проверка проекта и оформление документации',
    'Сдача зачета по практике',
]


def apply_font(run, size_pt):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size_pt)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_segments(paragraph, segments, size_pt=14):
    clear_paragraph(paragraph)
    for segment in segments:
        text = segment.get('text', '')
        if not text:
            continue
        run = paragraph.add_run(text)
        apply_font(run, size_pt)
        if segment.get('bold') is not None:
            run.bold = segment['bold']
        if segment.get('italic') is not None:
            run.italic = segment['italic']
        if segment.get('break_after'):
            run.add_break()


def set_plain_paragraph(paragraph, text, size_pt=14):
    lines = text.split('\n')
    segments = []
    for idx, line in enumerate(lines):
        segments.append({'text': line, 'break_after': idx < len(lines) - 1})
    set_paragraph_segments(paragraph, segments, size_pt=size_pt)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    apply_font(run, 12)


def build_diary():
    document = Document(str(TEMPLATE_PATH))

    set_plain_paragraph(
        document.paragraphs[11],
        f'{STUDENT_FULL_NAME}          гр.{GROUP}                   _______________',
    )
    set_plain_paragraph(document.paragraphs[14], START_DATE)
    set_plain_paragraph(document.paragraphs[15], END_DATE)
    set_plain_paragraph(
        document.paragraphs[19],
        f'преподаватель кафедры САИТ\n{SUPERVISOR_SHORT}            \t\t\t\t\t ________________',
    )
    set_plain_paragraph(
        document.paragraphs[24],
        f'преподаватель кафедры САИТ {SUPERVISOR_SHORT}         __________________',
    )
    set_paragraph_segments(
        document.paragraphs[29],
        [
            {'text': 'ФИО обучающегося, группа', 'bold': True},
            {'text': f' {STUDENT_FULL_NAME}, группа {GROUP}'},
        ],
    )
    set_paragraph_segments(
        document.paragraphs[30],
        [
            {'text': 'Место прохождения практики:', 'bold': True},
            {'text': f' {PRACTICE_PLACE}'},
        ],
    )

    table = document.tables[0]
    for row_index, task in enumerate(DAILY_TASKS, start=1):
        set_cell_text(table.rows[row_index].cells[1], task)

    document.save(str(OUTPUT_PATH))


if __name__ == '__main__':
    build_diary()
