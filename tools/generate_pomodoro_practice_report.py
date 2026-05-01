from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TEMPLATE_PATH = Path("/Users/kamilhus/Downloads/Отчет студента ФИИТ ПТП 231 Мудрик.docx")
OUTPUT_DIR = Path("/Users/kamilhus/Downloads/Хусаинов_Камиль_Фанисович")
OUTPUT_PATH = OUTPUT_DIR / "Отчет ПТП Pomodoro Web Хусаинов - итог.docx"

SCREENSHOT_DIR = Path(
    "/Users/kamilhus/Documents/Playground/tmp/showcase_screenshots/cropped"
)

STUDENT_FULL_NAME = "Хусаинов Камиль Фанисович"
GROUP = "09-231"
PRACTICE_YEAR = "2026"
PRACTICE_PERIOD = "с 10 февраля 2026 года по 25 марта 2026 года"
SCIENTIFIC_SUPERVISOR = "доцент кафедры САИТ, канд. техн. наук Курбангалеев А. А."
DEPARTMENT_SUPERVISOR = "ст. преподаватель КСАИТ Тихонова О. О."
ACCESS_DATE = "17.03.2026"

INTRO_TASKS = [
    "провести анализ предметной области и существующих сервисов личной продуктивности;",
    "сформулировать требования к веб-приложению Pomodoro Web с учетом сценариев ВКР;",
    "спроектировать архитектуру системы, базу данных и серверный API;",
    "реализовать приложение, выполнить первичное тестирование и подготовить локальный запуск.",
]

FUNCTIONAL_REQUIREMENTS = [
    "регистрация, вход и работа каждого пользователя только со своими данными;",
    "создание целей, задач, фокус-сессий и ежедневных фото-отчетов;",
    "автоматическая AI-проверка отчета с выдачей вердикта и пояснения;",
    "формирование мотивационной ленты, чата «Мотиватор» и статистики по цели.",
]

NON_FUNCTIONAL_REQUIREMENTS = [
    "безопасная аутентификация на основе пары JWT access/refresh токенов [8];",
    "REST-взаимодействие клиента и сервера по HTTP с JSON-обменом и multipart/form-data для изображений [7-9];",
    "запуск базовой конфигурации проекта одной командой Docker Compose [7];",
    "корректная работа интерфейса на настольных и мобильных устройствах и наличие автоматических тестов [5].",
]

SCENARIOS = [
    "постановка долгосрочной цели и формирование задач на текущий день;",
    "запуск и завершение фокус-сессии с фиксацией затраченного времени;",
    "загрузка фото-отчета с текстовым комментарием и получение AI-вердикта;",
    "получение мотивационных материалов и совета от «Мотиватора».",
]

LITERATURE = [
    "Pomofocus [Электронный ресурс]. URL: https://pomofocus.io/ (дата обращения: 17.03.2026).",
    "Focus To-Do [Электронный ресурс]. URL: https://www.focustodo.cn/ (дата обращения: 17.03.2026).",
    "TickTick [Электронный ресурс]. URL: https://ticktick.com/home (дата обращения: 17.03.2026).",
    "Spring Boot Documentation [Электронный ресурс]. URL: https://docs.spring.io/spring-boot/documentation.html (дата обращения: 17.03.2026).",
    "React Documentation [Электронный ресурс]. URL: https://react.dev/ (дата обращения: 17.03.2026).",
    "PostgreSQL Documentation [Электронный ресурс]. URL: https://www.postgresql.org/docs/ (дата обращения: 17.03.2026).",
    "Docker Compose Documentation [Электронный ресурс]. URL: https://docs.docker.com/compose/ (дата обращения: 17.03.2026).",
    "Introduction to JSON Web Tokens [Электронный ресурс]. URL: https://jwt.io/introduction (дата обращения: 17.03.2026).",
    "OpenAPI Specification [Электронный ресурс]. URL: https://swagger.io/specification/ (дата обращения: 17.03.2026).",
    "OpenAI API Platform Documentation [Электронный ресурс]. URL: https://platform.openai.com/docs/ (дата обращения: 17.03.2026).",
    "Ollama [Электронный ресурс]. URL: https://ollama.com/ (дата обращения: 17.03.2026).",
    "MediaWiki API: Search [Электронный ресурс]. URL: https://www.mediawiki.org/wiki/API:Search (дата обращения: 17.03.2026).",
]

def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def normalize_template_styles(document: Document) -> None:
    for style in document.styles:
        style_name = (style.name or "").lower()
        if style_name.startswith("toc "):
            try:
                r_pr = style._element.rPr
                if r_pr is None:
                    r_pr = OxmlElement("w:rPr")
                    style._element.append(r_pr)
                for tag in ("w:caps", "w:smallCaps"):
                    for node in list(r_pr.findall(qn(tag))):
                        r_pr.remove(node)
                caps = OxmlElement("w:caps")
                caps.set(qn("w:val"), "0")
                r_pr.append(caps)
                small_caps = OxmlElement("w:smallCaps")
                small_caps.set(qn("w:val"), "0")
                r_pr.append(small_caps)
            except AttributeError:
                continue


def apply_run_font(
    run, *, size_pt=14, bold=None, italic=None, font_name="Times New Roman"
):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_body_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def add_paragraph(
    document,
    text="",
    *,
    style="Normal",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=True,
    bold=False,
    italic=False,
    size_pt=14,
):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    set_body_format(paragraph, align=align, first_line=first_line)
    return paragraph


def add_heading(document, text, *, centered=False):
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=14, bold=True)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    return paragraph


def add_subheading(document, text):
    paragraph = document.add_paragraph(style="Heading 2")
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=14, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    return paragraph


def add_list(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Paragraph")
        run = paragraph.add_run(item)
        apply_run_font(run, size_pt=14)
        set_body_format(
            paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True
        )


def add_numbered_references(document, items):
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(f"{index}. {item}")
        apply_run_font(run, size_pt=14)
        set_body_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)


def add_blank(document, count=1):
    for _ in range(count):
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run("")
        apply_run_font(run, size_pt=14)
        set_body_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)


def add_page_break(document):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(document):
    heading = document.add_paragraph(style="Normal")
    run = heading.add_run("ОГЛАВЛЕНИЕ")
    apply_run_font(run, size_pt=14, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading.paragraph_format.first_line_indent = Cm(0)

    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    apply_run_font(run_instr, size_pt=14)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run("Оглавление будет обновлено в Microsoft Word")
    apply_run_font(run_text, size_pt=14)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_page_number_footer(document):
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        run_begin = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_begin)

        run_instr = paragraph.add_run()
        apply_run_font(run_instr, size_pt=12)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)

        run_sep = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)

        run_text = paragraph.add_run("1")
        apply_run_font(run_text, size_pt=12)

        run_end = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)


def add_code_block(document, title, code_text):
    add_paragraph(document, title, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    for line in code_text.splitlines():
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(line)
        apply_run_font(run, size_pt=10, font_name="Courier New")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)


def add_image(document, image_path: Path, caption: str, *, width_cm=13.4):
    if not image_path.exists():
        return
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_paragraph(document, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def build_title_page(document):
    title_lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "КАЗАНСКИЙ (ПРИВОЛЖСКИЙ) ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ",
        "Институт вычислительной математики и информационных технологий",
    ]
    for line in title_lines:
        add_paragraph(document, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_blank(document, 3)
    add_paragraph(document, "ОТЧЕТ", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(
        document,
        "по технологической (проектно-технологической) (производственной) практике",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    add_blank(document, 3)

    add_paragraph(
        document,
        f"Обучающийся: {STUDENT_FULL_NAME}, гр. {GROUP} ____________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        f"Научный руководитель: {SCIENTIFIC_SUPERVISOR} ____________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        f"Руководитель практики от кафедры: {DEPARTMENT_SUPERVISOR} ____________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "Оценка за практику ______________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        f"Дата сдачи отчета ______________________ {PRACTICE_YEAR} г.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 8)
    add_paragraph(
        document, f"Казань – {PRACTICE_YEAR}", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False
    )


def add_content(document):
    add_page_break(document)
    add_toc(document)
    add_page_break(document)

    add_heading(document, "ВВЕДЕНИЕ", centered=True)
    add_paragraph(
        document,
        f"Производственная практика проходила на кафедре системного анализа и информационных технологий Института вычислительной математики и информационных технологий КФУ {PRACTICE_PERIOD}.",
    )
    add_paragraph(
        document,
        "Целью практики являлись анализ предметной области, проектирование, реализация и первичное тестирование веб-приложения Pomodoro Web, разрабатываемого в рамках ВКР как единая среда для управления целями, задачами дня, фокус-сессиями и мотивационной поддержкой пользователя.",
    )
    add_paragraph(document, "В ходе практики решались следующие задачи:")
    add_list(document, INTRO_TASKS)

    add_heading(document, "1. Анализ предметной области и существующих решений")
    add_subheading(document, "1.1. Общая характеристика предметной области")
    add_paragraph(
        document,
        "Для пользователя, работающего над учебной или проектной целью, важны не только список задач и таймер, но и сохранение контекста: понимание текущего шага, учет затраченного времени и фиксация фактического результата. На практике эти функции часто распределены между несколькими сервисами, что увеличивает число переключений внимания. Поэтому в проекте особое внимание уделено фото-отчету, AI-проверке и ежедневной сводке прогресса.",
    )

    add_subheading(document, "1.2. Анализ существующих решений")
    add_paragraph(
        document,
        "В качестве ближайших аналогов были рассмотрены Pomofocus, Focus To-Do и TickTick [1-3]. Эти продукты поддерживают сценарии планирования и фокус-таймера, однако не объединяют в одном компактном интерфейсе долгосрочную цель, задачи дня, фото-подтверждение результата, мотивационную ленту и контекстный текстовый помощник. Именно эта совокупность сценариев была положена в основу Pomodoro Web.",
    )

    add_subheading(document, "1.3. Потребности целевой аудитории")
    add_paragraph(
        document,
        "Целевой аудиторией приложения являются студенты, начинающие разработчики и специалисты, которым требуется ежедневный контроль собственной дисциплины. Для данной аудитории характерны большое число отвлекающих факторов, работа короткими итерациями и необходимость быстро видеть накопленный прогресс.",
    )
    add_paragraph(
        document,
        "На основе анализа пользовательских сценариев были выделены следующие ключевые потребности:"
    )
    add_list(document, SCENARIOS)

    add_heading(document, "2. Требования к системе")
    add_subheading(document, "2.1. Функциональные требования")
    add_paragraph(
        document,
        "На основании предметного анализа и сценариев использования были сформулированы следующие функциональные требования к Pomodoro Web:"
    )
    add_list(document, FUNCTIONAL_REQUIREMENTS)

    add_subheading(document, "2.2. Нефункциональные требования")
    add_paragraph(
        document,
        "Нефункциональные требования определялись выбранным стеком, режимом локального запуска и требованиями к безопасности и сопровождаемости системы:"
    )
    add_list(document, NON_FUNCTIONAL_REQUIREMENTS)

    add_heading(document, "3. Проектирование веб-приложения Pomodoro Web")
    add_subheading(document, "3.1. Архитектура приложения")
    add_paragraph(
        document,
        "Система реализована как монорепозиторий с backend-модулем на Spring Boot, frontend-модулем на React и файлами локального развертывания [4, 5, 7]. Архитектура приложения является клиент-серверной: браузерный клиент отвечает за интерфейс и маршрутизацию, а серверный модуль выполняет бизнес-логику, аутентификацию, работу с данными, AI-вызовы и фоновые задачи. Основным хранилищем данных выбран PostgreSQL, а для изображений используется файловое хранилище backend с публичными путями /uploads/... [6].",
    )

    add_subheading(document, "3.2. Функциональные модули")
    add_paragraph(
        document,
        "В проекте выделены модули аутентификации, целей, задач, фокус-сессий, отчетов, мотивационной ленты, чата «Мотиватор» и статистики. Такое разбиение согласуется со структурой REST API и набором серверных сервисов: AuthService, GoalService, FocusSessionService, ReportService, MotivationService и ChatService.",
    )

    add_subheading(document, "3.3. Проектирование базы данных")
    add_paragraph(
        document,
        "Логическая схема данных задана миграциями Flyway и включает таблицы users, goals, tasks, focus_sessions, reports, motivation_images, motivation_quotes, chat_threads, chat_messages, refresh_tokens и daily_summaries. Особое внимание уделено дневным операциям: daily_summaries хранит агрегированный результат по дате, а refresh_tokens позволяет отзывать и очищать токены независимо от основной пользовательской записи [4, 6].",
    )

    add_subheading(document, "3.4. Проектирование API")
    add_paragraph(
        document,
        "Серверный обмен построен по REST-подходу с использованием HTTP и JSON [7, 9]. Базовый префикс API — /api. Для защищенных запросов применяется заголовок Authorization: Bearer <access-token>, а для обновления сессии используется маршрут /api/auth/refresh и refresh-токен [8]. Большинство операций используют application/json, а загрузка отчета выполняется через multipart/form-data.",
    )

    add_heading(document, "4. Реализация приложения")
    add_subheading(document, "4.1. Реализация серверной части")
    add_paragraph(
        document,
        "Backend реализован на Java 21 и Spring Boot 3 с использованием Spring Web, Spring Security, Spring Data JPA, Bean Validation и встроенного планировщика [4]. Регистрация и вход формируют пару JWT-токенов, при этом refresh-токен дополнительно сохраняется в таблице refresh_tokens, а истекшие записи очищаются при выдаче новой пары [8].",
    )
    add_paragraph(
        document,
        "Алгоритмы основных операций реализованы на уровне сервисов. При старте фокус-сессии система проверяет наличие задач на текущий день и отсутствие активной сессии; при остановке длительность вычисляется как разность между endedAt и startedAt. Планировщик DailyScheduler ежедневно переводит просроченные отчеты в OVERDUE, пересчитывает streak и формирует запись daily_summaries, а также обновляет мотивационные изображения и цитаты.",
    )

    add_subheading(document, "4.2. Реализация AI-интеграции")
    add_paragraph(
        document,
        "AI-подсистема инкапсулирована через интерфейс AiService и поддерживает три режима работы: mock, openai и local. В режиме openai анализ фото выполняется запросом к OpenAI API, где изображение передается как base64, а в контекст включаются цель, комментарий пользователя и список задач дня. От модели требуется вернуть JSON с полями verdict, confidence и explanation [10].",
    )
    add_paragraph(
        document,
        "В режимах mock и local реализован эвристический алгоритм: из контекста извлекаются задачи дня, из них выделяются ключевые слова, которые сравниваются с нормализованным комментарием пользователя. Дополнительно проверяются маркеры завершения, например «выполнил» и «готово». Для локального чата используется Ollama [11], а для мотивационной ленты — интернет-поиск изображений через MediaWiki API [12].",
    )

    add_subheading(document, "4.3. Реализация клиентской части")
    add_paragraph(
        document,
        "Frontend построен как SPA на React, TypeScript и Vite [5]. Маршруты /login и /register доступны без авторизации, а остальные страницы защищены компонентом ProtectedRoute. Модуль apiClient автоматически подставляет access-токен в заголовок Authorization, а при ответе 401 пытается обновить токены через /auth/refresh. Выбор активной цели и состояние навигации сохраняются в localStorage, а цвет цели задает CSS-переменные интерфейса.",
    )

    add_subheading(document, "4.4. Интеграция и локальный запуск")
    add_paragraph(
        document,
        "Интеграция проекта выполнена через Docker Compose [7]. Базовая конфигурация поднимает PostgreSQL, backend и frontend, а для локального LLM-режима может дополнительно использоваться профиль с Ollama [11]. Все основные параметры, включая порты, строку подключения к БД, режим AI и директорию uploads, задаются через переменные окружения. Это позволяет воспроизводимо запускать проект как на машине разработчика, так и в демонстрационной среде.",
    )

    add_heading(document, "5. Пользовательский интерфейс")
    add_subheading(document, "5.1. Основные экраны системы")
    add_paragraph(
        document,
        "Пользовательский интерфейс организован вокруг выбранной цели. Ключевым рабочим экраном является вкладка «Фокус», где пользователь создает задачи дня, запускает таймер и отправляет фото-отчет. Остальные экраны предоставляют обзор прогресса, мотивационные материалы и статистику.",
    )
    add_paragraph(
        document,
        "После локального запуска проекта были вручную проверены основные пользовательские экраны: главная панель, вкладка «Контроль», рабочий экран «Фокус», мотивационная лента, чат и статистика. Наиболее важным с точки зрения ежедневного сценария является экран «Фокус», где объединены задачи дня, таймер и загрузка отчета.",
    )

    add_heading(document, "6. Первичное тестирование")
    add_paragraph(
        document,
        "Первичное тестирование проекта включало автоматические и ручные проверки. В backend-модуле реализованы 4 интеграционных тестовых класса с 13 сценариями; при подготовке отчета этот набор был успешно выполнен в контейнере Maven (`EXIT:0`). Frontend-модуль содержит 4 тестовых файла и 7 unit-тестов; при запуске в контейнере Node 22 тесты завершились успешно, после чего была выполнена production-сборка без ошибок. Ручная проверка после запуска Docker Compose подтвердила корректность основных пользовательских сценариев.",
    )

    add_heading(document, "ЗАКЛЮЧЕНИЕ", centered=True)
    add_paragraph(
        document,
        "В ходе производственной практики был разработан и протестирован MVP веб-приложения Pomodoro Web, объединяющего постановку цели, контроль задач дня, фокус-сессии, фото-подтверждение результата, мотивационную ленту и AI-чат. Выполненные на практике аналитические, проектные и программные решения могут быть непосредственно использованы при подготовке выпускной квалификационной работы.",
    )

    add_heading(document, "СПИСОК ЛИТЕРАТУРЫ", centered=True)
    add_numbered_references(document, LITERATURE)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(str(TEMPLATE_PATH))
    normalize_template_styles(document)
    clear_document(document)
    add_page_number_footer(document)
    build_title_page(document)
    add_content(document)
    document.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
