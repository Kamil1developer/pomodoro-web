from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPORT_PATH = Path("/Users/kamilhus/Downloads/Отчет ЭП Pomodoro Web Хусаинов.docx")
BACKUP_PATH = REPORT_PATH.with_name(REPORT_PATH.stem + " - backup before normalize.docx")

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, font_name: str, size_pt: float, *, bold=None) -> None:
    run.font.name = font_name
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def is_code_paragraph(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False

    run_fonts = {run.font.name for run in paragraph.runs if run.text.strip()}
    if CODE_FONT in run_fonts:
        return True

    code_prefixes = (
        "docker ",
        "services:",
        "postgres:",
        "backend:",
        "frontend:",
        "image:",
        "build:",
        "public class ",
        "public interface ",
        "@Entity",
        "@Table",
        "private ",
        "return ",
        "context:",
        "environment:",
        "ports:",
    )
    return text.startswith(code_prefixes)


def normalize_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    style_name = paragraph.style.name or ""

    if is_code_paragraph(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, CODE_FONT, 12, bold=False)
        return

    if style_name.startswith("Heading"):
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        if style_name == "Heading 1":
            if text in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ЛИТЕРАТУРЫ", "ПРИЛОЖЕНИЯ"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(1.25)
        elif style_name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)

        for run in paragraph.runs:
            set_run_font(run, BODY_FONT, 14, bold=True)
        return

    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    centered_titles = {
        "ОГЛАВЛЕНИЕ",
        "ОТЧЕТ",
        "по эксплуатационной (производственной) практике",
        "Казань – 2026",
    }
    left_labels = {
        "Научный руководитель",
        "Руководитель практики от кафедры:",
        "Команда запуска проекта:",
        "Команда запуска backend-тестов:",
        "Команда запуска frontend-проверки:",
    }

    if text in centered_titles:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif text in left_labels or text.startswith("Таблица "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif text.startswith("Рисунок "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif style_name == "List Paragraph":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif text:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0)

    for run in paragraph.runs:
        # Only headings should remain bold.
        set_run_font(run, BODY_FONT, 14, bold=False)


def normalize_table(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif cell_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    set_run_font(run, BODY_FONT, 14, bold=(row_index == 0))


def main() -> None:
    if not REPORT_PATH.exists():
        raise SystemExit(f"File not found: {REPORT_PATH}")

    shutil.copy2(REPORT_PATH, BACKUP_PATH)
    document = Document(str(REPORT_PATH))

    for section in document.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = BLACK

    for paragraph in document.paragraphs:
        normalize_paragraph(paragraph)

    for table in document.tables:
        normalize_table(table)

    document.save(str(REPORT_PATH))
    print(f"backup: {BACKUP_PATH}")
    print(f"updated: {REPORT_PATH}")


if __name__ == "__main__":
    main()
