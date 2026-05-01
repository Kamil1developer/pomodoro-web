from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT_PATH = Path("/Users/kamilhus/Downloads/Отчет ЭП Pomodoro Web Хусаинов.docx")
BACKUP_PATH = REPORT_PATH.with_name(REPORT_PATH.stem + " - backup before cleanup.docx")
IMAGE_WIDTH_CM = 14.0


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def set_toc_field(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    run_instr.font.name = "Times New Roman"
    run_instr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run_instr.font.size = Pt(14)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run("Оглавление будет обновлено в Microsoft Word")
    run_text.font.name = "Times New Roman"
    run_text._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run_text.font.size = Pt(14)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def normalize_images(document: Document) -> None:
    new_width = Cm(IMAGE_WIDTH_CM)
    for shape in document.inline_shapes:
        old_width = shape.width
        old_height = shape.height
        if old_width:
            ratio = float(old_height) / float(old_width)
            shape.width = new_width
            shape.height = int(new_width * ratio)


def main() -> None:
    if not REPORT_PATH.exists():
        raise SystemExit(f"Report not found: {REPORT_PATH}")

    shutil.copy2(REPORT_PATH, BACKUP_PATH)
    document = Document(str(REPORT_PATH))

    appendix_start = None
    toc_paragraph = None

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text == "ОГЛАВЛЕНИЕ":
            toc_paragraph = document.paragraphs[index + 1]
        elif text == "ПРИЛОЖЕНИЯ":
            appendix_start = index
            break

    if appendix_start is not None:
        for paragraph in list(document.paragraphs[appendix_start:]):
            remove_paragraph(paragraph)

    if toc_paragraph is not None:
        set_toc_field(toc_paragraph)

    replacements = {
        "Казань – 2026": "Казань",
        "Эксплуатационная практика проходила на кафедре системного анализа и информационных технологий Института вычислительной математики и информационных технологий КФУ с 26 марта 2026 года по 22 апреля 2026 года. В рамках практики рассматривался собственный учебный проект Pomodoro Web — веб-приложение для управления целями, задачами дня, фокус-сессиями и мотивационной поддержкой пользователя.":
            "Эксплуатационная практика проходила на кафедре системного анализа и информационных технологий Института вычислительной математики и информационных технологий КФУ. В рамках практики рассматривался собственный учебный проект Pomodoro Web — веб-приложение для управления целями, задачами дня, фокус-сессиями и мотивационной поддержкой пользователя.",
        "Первичная эксплуатация проводилась 13 апреля 2026 года на локальной машине разработчика. Для запуска использовался исходный монорепозиторий проекта с модулями `frontend`, `backend` и файлом `docker-compose.yml`. Базовый запуск приложения выполнялся одной командой, что соответствует выбранной архитектуре контейнеризации [4].":
            "Первичная эксплуатация проводилась на локальной машине разработчика. Для запуска использовался исходный монорепозиторий проекта с модулями `frontend`, `backend` и файлом `docker-compose.yml`. Базовый запуск приложения выполнялся одной командой, что соответствует выбранной архитектуре контейнеризации [4].",
        "Для демонстрационной цели в рабочем экземпляре приложения были подготовлены данные, отображаемые на экранах интерфейса. В результате в системе корректно отображались 2 фото-отчета, 3 мотивационные карточки, 4 сообщения в истории чата, прогресс 2 из 3 задач, суммарно 95 минут фокус-сессий и streak, равный 5 дням. Также на экране мотивации была показана ежедневная цитата на дату 12.04.2026.":
            "Для демонстрационной цели в рабочем экземпляре приложения были подготовлены данные, отображаемые на экранах интерфейса. В результате в системе корректно отображались 2 фото-отчета, 3 мотивационные карточки, 4 сообщения в истории чата, прогресс 2 из 3 задач, суммарно 95 минут фокус-сессий и streak, равный 5 дням. Также на экране мотивации была показана ежедневная цитата.",
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            replace_paragraph_text(paragraph, replacements[text])
            continue

        if "дата обращения:" in paragraph.text:
            cleaned = re.sub(r"\s*\(дата обращения:[^)]+\)\.?", ".", paragraph.text)
            cleaned = cleaned.replace("..", ".").strip()
            replace_paragraph_text(paragraph, cleaned)

    normalize_images(document)
    document.save(str(REPORT_PATH))
    print(f"backup: {BACKUP_PATH}")
    print(f"updated: {REPORT_PATH}")


if __name__ == "__main__":
    main()
