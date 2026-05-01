from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TEMPLATE_PATH = Path(
    "/Users/kamilhus/Downloads/в4_Отчет_студента_ФИИТ_ПТП_231_Хусаинов.docx"
)
OUTPUT_DIR = Path("/Users/kamilhus/Downloads/Хусаинов_Камиль_Фанисович")
OUTPUT_PATH = OUTPUT_DIR / "Отчет ЭП Pomodoro Web Хусаинов - стиль в4.docx"
SCREENSHOT_DIR = Path("/Users/kamilhus/Documents/Playground/tmp/report_run_screenshots")

STUDENT_FULL_NAME = "Хусаинов Камиль Фанисович"
GROUP = "09-231"
PRACTICE_YEAR = "2026"
PRACTICE_PERIOD = "с 26 марта 2026 года по 22 апреля 2026 года"
SCIENTIFIC_SUPERVISOR = "доцент кафедры САИТ, канд. тех. наук Курбангалеев А. А."
DEPARTMENT_SUPERVISOR = "ст.преподаватель КСАИТ Тихонова О.О."
TEST_DATE = "13 апреля 2026 года"
ACCESS_DATE = "13.04.2026"

INTRO_TASKS = [
    "подготовить среду локального запуска Pomodoro Web и провести первичную эксплуатацию приложения;",
    "проверить корректность основных пользовательских сценариев на работающем экземпляре системы;",
    "выполнить автоматические backend- и frontend-тесты и зафиксировать их фактические результаты;",
    "подготовить отчет по эксплуатационной практике на основе реальных данных проекта.",
]

FUNCTIONS = [
    "регистрация и аутентификация пользователя с JWT access/refresh токенами;",
    "создание целей, задач и отслеживание прогресса по выбранной цели;",
    "запуск и завершение фокус-сессий Pomodoro с накоплением статистики;",
    "загрузка фото-отчетов и вывод AI-вердикта по выполненной работе;",
    "мотивационная лента, ежедневная цитата и чат «Мотиватор»;",
    "страница статистики с динамикой выполненных задач, минут фокуса и streak.",
]

MANUAL_SCENARIOS = [
    (
        "Запуск приложения",
        "Выполнена команда `docker compose up --build`, после чего frontend стал доступен на http://localhost:5173, а backend API на http://localhost:18080.",
        "успешно",
    ),
    (
        "Первичная навигация по UI",
        "Через Selenium получены актуальные скриншоты страниц входа, главной панели, «Контроль», «Фокус», «Мотивация», «Мотиватор» и «Статистика».",
        "успешно",
    ),
    (
        "Проверка прикладных данных",
        "Для демонстрационной цели отображаются 2 отчета, 3 мотивационные карточки, 4 сообщения чата и статистика за 7 дней.",
        "успешно",
    ),
    (
        "API-сценарий регистрации и работы с задачами",
        "Создан отдельный тестовый пользователь, затем через API создана цель, добавлены 2 задачи и выполнены start/stop запросы для фокус-сессии.",
        "успешно",
    ),
]

BACKEND_TESTS = [
    (
        "AuthIntegrationTest",
        3,
        "проверка регистрации, входа и обновления токенов",
    ),
    (
        "GoalTaskFocusIntegrationTest",
        5,
        "проверка целей, задач и фокус-сессий",
    ),
    (
        "ReportMotivationIntegrationTest",
        2,
        "проверка отчетов и мотивационного модуля",
    ),
    (
        "ChatSchedulerIntegrationTest",
        3,
        "проверка чата и планировщика DailyScheduler",
    ),
]

FRONTEND_TESTS = [
    ("src/lib/format.test.ts", 3, "форматирование минут и процентов"),
    ("src/lib/authStorage.test.ts", 2, "сохранение и очистка токенов"),
    ("src/components/ProgressCard.test.tsx", 1, "отображение карточки прогресса"),
    ("src/components/GoalSelector.test.tsx", 1, "выбор активной цели"),
]

LITERATURE = [
    f"Spring Boot Documentation [Электронный ресурс]. URL: https://docs.spring.io/spring-boot/documentation.html (дата обращения: {ACCESS_DATE}).",
    f"React Documentation [Электронный ресурс]. URL: https://react.dev/ (дата обращения: {ACCESS_DATE}).",
    f"PostgreSQL Documentation [Электронный ресурс]. URL: https://www.postgresql.org/docs/ (дата обращения: {ACCESS_DATE}).",
    f"Docker Compose Documentation [Электронный ресурс]. URL: https://docs.docker.com/compose/ (дата обращения: {ACCESS_DATE}).",
    f"Vite Guide [Электронный ресурс]. URL: https://vite.dev/guide/ (дата обращения: {ACCESS_DATE}).",
    f"Introduction to JSON Web Tokens [Электронный ресурс]. URL: https://jwt.io/introduction (дата обращения: {ACCESS_DATE}).",
    f"OpenAPI Specification [Электронный ресурс]. URL: https://swagger.io/specification/ (дата обращения: {ACCESS_DATE}).",
    f"Pomofocus [Электронный ресурс]. URL: https://pomofocus.io/ (дата обращения: {ACCESS_DATE}).",
]


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def normalize_template_styles(document: Document) -> None:
    for style in document.styles:
        style_name = (style.name or "").lower()
        if style_name.startswith("toc "):
            try:
                r_pr = style._element.rPr
                if r_pr is None:
                    r_pr = OxmlElement("w:rPr")
                    style._element.append(r_pr)
                for tag in ("w:caps", "w:smallCaps"):
                    for node in list(r_pr.findall(qn(tag))):
                        r_pr.remove(node)
                caps = OxmlElement("w:caps")
                caps.set(qn("w:val"), "0")
                r_pr.append(caps)
                small_caps = OxmlElement("w:smallCaps")
                small_caps.set(qn("w:val"), "0")
                r_pr.append(small_caps)
            except AttributeError:
                continue


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def apply_run_font(
    run, *, size_pt=14, bold=None, italic=None, font_name="Times New Roman"
):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_body_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def add_paragraph(
    document,
    text="",
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=True,
    bold=False,
    italic=False,
    size_pt=14,
):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    set_body_format(paragraph, align=align, first_line=first_line)
    return paragraph


def add_heading(document, text, *, centered=False):
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=14, bold=True)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(0 if centered else 1.25)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    return paragraph


def add_subheading(document, text):
    paragraph = document.add_paragraph(style="Heading 2")
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=14, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    return paragraph


def add_numbered_list(document, items):
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(f"{index}) {item}")
        apply_run_font(run, size_pt=14)
        set_body_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)


def add_blank(document, count=1):
    for _ in range(count):
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run("")
        apply_run_font(run, size_pt=14)
        set_body_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)


def add_page_break(document):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(document):
    heading = document.add_paragraph(style="Normal")
    run = heading.add_run("ОГЛАВЛЕНИЕ")
    apply_run_font(run, size_pt=14, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading.paragraph_format.first_line_indent = Cm(0)

    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    apply_run_font(run_instr, size_pt=14)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run("Оглавление будет обновлено в Microsoft Word")
    apply_run_font(run_text, size_pt=14)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_page_number_footer(document):
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        run_begin = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_begin)

        run_instr = paragraph.add_run()
        apply_run_font(run_instr, size_pt=12)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)

        run_sep = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)

        run_text = paragraph.add_run("1")
        apply_run_font(run_text, size_pt=12)

        run_end = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)


def add_code_block(document, title, code_text):
    add_paragraph(document, title, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    for line in code_text.splitlines():
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(line)
        apply_run_font(run, size_pt=12, font_name="Courier New")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=12):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    apply_run_font(run, size_pt=size_pt, bold=bold)


def add_results_table(document, headers, rows, widths_cm):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = None
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        table.columns[index].width = Cm(widths_cm[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], str(value), align=align)

    for paragraph in document.paragraphs[-1:]:
        paragraph.paragraph_format.space_after = Pt(0)
    add_blank(document, 1)
    return table


def add_image(document, image_path: Path, caption: str, *, width_cm=14.8):
    if not image_path.exists():
        return
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_paragraph(document, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_numbered_references(document, items):
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(f"•\t{item}")
        apply_run_font(run, size_pt=14)
        set_body_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)


def add_competence_table(document):
    headers = [
        "Компетенция",
        "Расшифровка",
        "Приобретенные знания, умения и навыки",
    ]
    rows = [
        (
            "УК-1",
            "Способен осуществлять поиск, критический анализ и синтез информации, применять системный подход для решения поставленных задач",
            "В ходе практики выполнен анализ предметной области, сопоставлены пользовательские сценарии и результаты тестирования, подготовлена структурированная сводка по проекту Pomodoro Web.",
        ),
        (
            "ПК-1",
            "Проверка работоспособности и рефакторинг кода программного обеспечения",
            "Проведены backend- и frontend-тесты, выполнен локальный запуск приложения и подтверждена корректность основных сценариев эксплуатации.",
        ),
        (
            "ПК-2",
            "Интеграция программных модулей и компонент и верификация выпусков программного продукта",
            "Проверена совместная работа React frontend, Spring Boot backend и PostgreSQL в окружении Docker Compose, а также успешность production-сборки клиентской части.",
        ),
        (
            "ПК-3",
            "Разработка требований и проектирование программного обеспечения",
            "В отчете систематизированы назначение системы, функциональные возможности и требования к эксплуатации разработанного приложения.",
        ),
        (
            "ПК-4",
            "Оценка и выбор варианта архитектуры программного средства",
            "На практике подтверждена работоспособность клиент-серверной архитектуры монорепозитория с отдельными модулями frontend, backend и инфраструктурой запуска.",
        ),
    ]
    add_results_table(document, headers, rows, [2.2, 6.3, 8.5])


def build_title_page(document):
    title_lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "КАЗАНСКИЙ (ПРИВОЛЖСКИЙ) ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ",
        "Институт вычислительной математики и информационных технологий",
    ]
    for line in title_lines:
        add_paragraph(document, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_blank(document, 3)
    add_paragraph(document, "ОТЧЕТ", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(
        document,
        "по эксплуатационной (производственной) практике",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    add_blank(document, 3)

    add_paragraph(
        document,
        f"Обучающийся {STUDENT_FULL_NAME} _гр.{GROUP}_    _________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_paragraph(
        document,
        "\t\t         (ФИО студента)                            (Группа)                              (Подпись)",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "Научный руководитель",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        f"{SCIENTIFIC_SUPERVISOR}\t\t\t\t__________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "(должность, степень ФИО)                                                                                         (Подпись)",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "Руководитель практики от кафедры:",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        f"{DEPARTMENT_SUPERVISOR}\t                      __________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "Оценка за практику ______________________              __________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 1)
    add_paragraph(
        document,
        "Дата сдачи отчета _______________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank(document, 8)
    add_paragraph(
        document,
        f"Казань – {PRACTICE_YEAR}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )


def add_content(document):
    add_page_break(document)
    add_toc(document)
    add_page_break(document)

    add_heading(document, "ВВЕДЕНИЕ", centered=True)
    add_paragraph(
        document,
        f"Эксплуатационная практика проходила на кафедре системного анализа и информационных технологий Института вычислительной математики и информационных технологий КФУ {PRACTICE_PERIOD}. В рамках практики рассматривался собственный учебный проект Pomodoro Web — веб-приложение для управления целями, задачами дня, фокус-сессиями и мотивационной поддержкой пользователя.",
    )
    add_paragraph(
        document,
        "Целью практики являлись первичная эксплуатация разработанного программного обеспечения, проверка его работоспособности в локальной среде и фиксация реальных результатов тестирования без использования недостоверных или неподтвержденных данных.",
    )
    add_paragraph(document, "В ходе практики были решены следующие задачи:")
    add_numbered_list(document, INTRO_TASKS)

    add_heading(document, "1. Краткая характеристика разработанного программного обеспечения")
    add_subheading(document, "1.1. Назначение проекта Pomodoro Web")
    add_paragraph(
        document,
        "Pomodoro Web представляет собой веб-приложение для персонального контроля продуктивности. Система объединяет в одном интерфейсе работу с целями, задачами, таймером фокус-сессий, ежедневными отчетами с изображением, мотивационной лентой и текстовым AI-помощником. По сравнению с узкоспециализированными таймерами, например Pomofocus [8], проект ориентирован не только на отсчет времени, но и на сохранение контекста достижения цели и фиксацию результата работы.",
    )

    add_subheading(document, "1.2. Реализованные функции")
    add_paragraph(
        document,
        "В рамках MVP-версии приложения к моменту подготовки отчета реализованы следующие основные функции:"
    )
    add_numbered_list(document, FUNCTIONS)

    add_subheading(document, "1.3. Технологическая основа системы")
    add_paragraph(
        document,
        "Серверная часть приложения реализована на Java 21 и Spring Boot 3 [1], клиентская часть — на React, TypeScript и Vite [2, 5]. Для хранения данных используется PostgreSQL [3], локальное развертывание организовано через Docker Compose [4]. Аутентификация построена на JWT access/refresh токенах [6], а описание HTTP API доступно через Swagger/OpenAPI [7]. Такое сочетание технологий обеспечивает воспроизводимый локальный запуск и разделение ответственности между frontend, backend и уровнем хранения данных.",
    )

    add_heading(document, "2. Первичная эксплуатация разработанного ПО")
    add_subheading(document, "2.1. Подготовка среды и запуск")
    add_paragraph(
        document,
        f"Первичная эксплуатация проводилась {TEST_DATE} на локальной машине разработчика. Для запуска использовался исходный монорепозиторий проекта с модулями `frontend`, `backend` и файлом `docker-compose.yml`. Базовый запуск приложения выполнялся одной командой, что соответствует выбранной архитектуре контейнеризации [4].",
    )
    add_code_block(
        document,
        "Команда запуска проекта:",
        "docker compose up --build",
    )
    add_paragraph(
        document,
        "После запуска был подтвержден доступ к пользовательскому интерфейсу по адресу http://localhost:5173 и к backend API по адресу http://localhost:18080. Дополнительно проверена доступность Swagger UI, что позволило убедиться в работоспособности слоя REST API.",
    )

    add_subheading(document, "2.2. Проверка сценариев эксплуатации")
    add_paragraph(
        document,
        "Для эксплуатационной проверки были выполнены как пользовательские сценарии через графический интерфейс, так и отдельный API-сценарий с созданием нового пользователя. Сводка проведенных действий приведена в таблице 1."
    )
    add_paragraph(document, "Таблица 1 – Сценарии первичной эксплуатации", first_line=False)
    add_results_table(
        document,
        ["№", "Проверяемый сценарий", "Фактический результат", "Статус"],
        [
            (index, title, result, status)
            for index, (title, result, status) in enumerate(MANUAL_SCENARIOS, start=1)
        ],
        [1.0, 4.4, 8.6, 2.0],
    )
    add_paragraph(
        document,
        "Для демонстрационной цели в рабочем экземпляре приложения были подготовлены данные, отображаемые на экранах интерфейса. В результате в системе корректно отображались 2 фото-отчета, 3 мотивационные карточки, 4 сообщения в истории чата, прогресс 2 из 3 задач, суммарно 95 минут фокус-сессий и streak, равный 5 дням. Также на экране мотивации была показана ежедневная цитата на дату 12.04.2026.",
    )
    add_paragraph(
        document,
        "Отдельно был проведен API-сценарий без предварительно подготовленных данных: зарегистрирован новый пользователь, создана тестовая цель, добавлены 2 задачи и выполнены запросы на запуск и завершение фокус-сессии. Полученные ответы сервера подтвердили, что сценарий регистрации, работы с задачами и фиксации фокус-сессии проходит без ошибок.",
    )

    add_subheading(document, "2.3. Визуальная проверка интерфейса")
    add_paragraph(
        document,
        "Для подтверждения первичной эксплуатации были сняты свежие скриншоты с работающего приложения. На рисунках ниже приведены наиболее показательные экраны, проверенные в ходе практики.",
    )
    add_image(
        document,
        SCREENSHOT_DIR / "02-dashboard.png",
        "Рисунок 1 – Главная панель Pomodoro Web с последними отчетами и сводкой по целям.",
    )
    add_image(
        document,
        SCREENSHOT_DIR / "04-focus.png",
        "Рисунок 2 – Экран «Фокус» с задачами дня, таймером и блоком фото-отчета.",
    )
    add_image(
        document,
        SCREENSHOT_DIR / "07-stats.png",
        "Рисунок 3 – Экран статистики с данными по выполненным задачам и минутам фокуса.",
    )

    add_heading(document, "3. Тестирование разработанного ПО")
    add_subheading(document, "3.1. Автоматические backend-тесты")
    add_paragraph(
        document,
        "Для серверной части был выполнен реальный прогон интеграционных тестов в контейнере Maven. Тестовый набор использует Spring Boot Test и проверяет аутентификацию, цели, задачи, фокус-сессии, отчеты, мотивационный модуль, чат и планировщик [1, 3]. Запуск производился следующей командой:"
    )
    add_code_block(
        document,
        "Команда запуска backend-тестов:",
        "docker run --rm -v \"$PWD/backend\":/app -w /app maven:3.9.9-eclipse-temurin-21 mvn test",
    )
    add_paragraph(
        document,
        "По итогам выполнения backend-тестов получен результат `BUILD SUCCESS`, завершение с кодом `0`, суммарно выполнено 13 тестов без ошибок и пропусков. Детализация набора приведена в таблице 2."
    )
    add_paragraph(document, "Таблица 2 – Результаты backend-тестов", first_line=False)
    add_results_table(
        document,
        ["№", "Тестовый класс", "Количество тестов", "Назначение"],
        [
            (index, name, count, description)
            for index, (name, count, description) in enumerate(BACKEND_TESTS, start=1)
        ],
        [1.0, 5.8, 3.0, 7.2],
    )

    add_subheading(document, "3.2. Автоматические frontend-тесты и production-сборка")
    add_paragraph(
        document,
        "Для клиентской части был выполнен полный прогон unit-тестов и затем production-сборка приложения в контейнере Node 22. Такой подход позволил проверить как корректность прикладных функций, так и способность проекта собраться в production-режиме [2, 5].",
    )
    add_code_block(
        document,
        "Команда запуска frontend-проверки:",
        "docker run --rm -v \"$PWD/frontend\":/app -w /app node:22-alpine sh -lc 'npm install && npm run test:run && npm run build'",
    )
    add_paragraph(
        document,
        "Результат выполнения: 4 тестовых файла и 7 тестов завершились успешно, после чего production-сборка `vite build` также завершилась без ошибок. Итоговая сборка сформировала клиентские ресурсы в каталоге `dist/`.",
    )
    add_paragraph(document, "Таблица 3 – Результаты frontend-тестов", first_line=False)
    add_results_table(
        document,
        ["№", "Файл тестов", "Количество тестов", "Проверяемая область"],
        [
            (index, name, count, description)
            for index, (name, count, description) in enumerate(FRONTEND_TESTS, start=1)
        ],
        [1.0, 6.1, 3.0, 6.9],
    )

    add_subheading(document, "3.3. Итоги тестирования и замечания")
    add_paragraph(
        document,
        "Таким образом, при подготовке отчета были реально выполнены 20 автоматических тестов: 13 на backend и 7 на frontend. Все тесты завершились успешно. Дополнительно подтверждена работоспособность ручных сценариев эксплуатации и доступность основных экранов приложения в браузере.",
    )
    add_paragraph(
        document,
        "В процессе frontend-проверки менеджер пакетов `npm` сообщил о наличии 7 уязвимостей зависимостей (5 moderate и 2 high). Эти предупреждения не помешали прохождению тестов и production-сборке, однако должны быть учтены как отдельная задача по сопровождению проекта.",
    )

    add_heading(document, "ЗАКЛЮЧЕНИЕ", centered=True)
    add_paragraph(
        document,
        "В ходе эксплуатационной практики был проведен запуск, первичная эксплуатация и тестирование веб-приложения Pomodoro Web. Проверка показала, что система корректно разворачивается в локальной среде, предоставляет заявленные пользовательские функции и проходит автоматические проверки без ошибок.",
    )
    add_paragraph(
        document,
        "Главным результатом практики стало не только подтверждение работоспособности проекта, но и фиксация фактических данных эксплуатации: выполненных тестов, доступных экранов интерфейса, сценариев использования и выявленных замечаний по зависимостям. Полученные результаты могут быть использованы как практическая основа для дальнейшего развития проекта и подготовки выпускной квалификационной работы.",
    )
    add_paragraph(
        document,
        "За период практики были приобретены следующие компетенции:",
    )
    add_paragraph(document, "Таблица 4 – Компетенции", first_line=False)
    add_competence_table(document)

    add_heading(document, "СПИСОК ЛИТЕРАТУРЫ", centered=True)
    add_numbered_references(document, LITERATURE)

    add_heading(document, "ПРИЛОЖЕНИЯ", centered=True)
    add_paragraph(
        document,
        "Приложение А. Дополнительные скриншоты интерфейса, полученные при первичной эксплуатации проекта.",
        first_line=False,
    )
    add_image(
        document,
        SCREENSHOT_DIR / "03-control.png",
        "Рисунок А.1 – Экран «Контроль» с редактированием цели и прогрессом.",
    )
    add_image(
        document,
        SCREENSHOT_DIR / "05-motivation.png",
        "Рисунок А.2 – Экран «Мотивация» с карточками и ежедневной цитатой.",
    )
    add_image(
        document,
        SCREENSHOT_DIR / "06-chat.png",
        "Рисунок А.3 – Экран «Мотиватор» с сохраненной историей диалога.",
    )


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(str(TEMPLATE_PATH))
    normalize_template_styles(document)
    clear_document(document)
    build_title_page(document)
    add_content(document)
    document.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
